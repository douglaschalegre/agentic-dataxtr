"""Document parsing service for multiple formats."""

import base64
import re
from contextvars import ContextVar
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Literal, Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image

# Context variable to store current parser instance
_current_parser: ContextVar[Optional["DocumentParser"]] = ContextVar(
    "current_parser", default=None
)


@dataclass
class Section:
    """Document section with header and content."""

    title: str
    content: str
    page_number: int
    level: int = 1


@dataclass
class SearchMatch:
    """Result from searching document content."""

    text: str
    page: int
    position: int
    confidence: float = 1.0


@dataclass
class DocumentParser:
    """Unified document parser supporting multiple formats."""

    file_path: Path
    document_type: Literal["pdf", "image", "docx", "xlsx", "csv"]

    # Parsed content
    _pages: list[str] = field(default_factory=list)
    _sections: list[Section] = field(default_factory=list)
    _images: dict[int, list[bytes]] = field(default_factory=dict)
    _tables: dict[int, list[dict]] = field(default_factory=dict)

    # Metadata
    page_count: int = 0
    has_tables: bool = False
    has_images: bool = False

    @classmethod
    def get_current(cls) -> "DocumentParser":
        """Get the current parser from context."""
        parser = _current_parser.get()
        if parser is None:
            raise RuntimeError("No document parser in context")
        return parser

    def set_current(self) -> None:
        """Set this parser as the current context."""
        _current_parser.set(self)

    async def load(self) -> tuple[dict, dict]:
        """Load and parse the document.

        Returns:
            Tuple of (content dict, metadata dict)
        """
        if self.document_type == "pdf":
            return await self._load_pdf()
        elif self.document_type == "image":
            return await self._load_image()
        elif self.document_type == "docx":
            return await self._load_docx()
        elif self.document_type == "xlsx":
            return await self._load_xlsx()
        elif self.document_type == "csv":
            return await self._load_csv()
        else:
            raise ValueError(f"Unsupported document type: {self.document_type}")

    async def _load_pdf(self) -> tuple[dict, dict]:
        """Load and parse a PDF document."""
        doc = fitz.open(str(self.file_path))
        self.page_count = len(doc)

        full_text = []
        for page_num, page in enumerate(doc):
            # Extract text
            text = page.get_text()
            self._pages.append(text)
            full_text.append(text)

            # Check for images
            image_list = page.get_images()
            if image_list:
                self.has_images = True
                self._images[page_num] = []
                for img in image_list:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    self._images[page_num].append(base_image["image"])

            # Check for tables (simple heuristic)
            if "|" in text or "\t" in text:
                self.has_tables = True

        # Extract sections based on font size changes
        self._sections = self._extract_sections_from_pdf(doc)

        doc.close()

        content = {
            "full_text": "\n\n".join(full_text),
            "pages": self._pages,
            "sections": [
                {"title": s.title, "content": s.content, "page": s.page_number}
                for s in self._sections
            ],
        }

        metadata = {
            "page_count": self.page_count,
            "has_tables": self.has_tables,
            "has_images": self.has_images,
            "file_size": self.file_path.stat().st_size,
        }

        return content, metadata

    def _extract_sections_from_pdf(self, doc: fitz.Document) -> list[Section]:
        """Extract sections based on formatting."""
        sections = []
        current_section = None

        for page_num, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "lines" not in block:
                    continue

                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        font_size = span["size"]

                        # Heuristic: larger font = section header
                        if font_size > 14 and len(text) < 100:
                            if current_section:
                                sections.append(current_section)
                            current_section = Section(
                                title=text,
                                content="",
                                page_number=page_num + 1,
                                level=1 if font_size > 18 else 2,
                            )
                        elif current_section:
                            current_section.content += text + " "

        if current_section:
            sections.append(current_section)

        return sections

    async def _load_image(self) -> tuple[dict, dict]:
        """Load an image document."""
        self.page_count = 1
        self.has_images = True

        # Read image and encode as base64 for vision models
        with open(self.file_path, "rb") as f:
            image_data = f.read()

        self._images[0] = [image_data]

        content = {
            "full_text": "",  # Will be populated by OCR if needed
            "pages": [""],
            "image_base64": base64.b64encode(image_data).decode(),
        }

        img = Image.open(self.file_path)
        metadata = {
            "page_count": 1,
            "has_tables": False,
            "has_images": True,
            "image_size": img.size,
            "file_size": self.file_path.stat().st_size,
        }

        return content, metadata

    async def _load_docx(self) -> tuple[dict, dict]:
        """Load a Word document."""
        doc = DocxDocument(str(self.file_path))

        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Extract tables
        if doc.tables:
            self.has_tables = True
            self._tables[0] = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                self._tables[0].append({"rows": rows})

        self.page_count = 1  # DOCX doesn't have page concept until rendered
        self._pages = ["\n".join(full_text)]

        content = {
            "full_text": "\n".join(full_text),
            "pages": self._pages,
            "tables": self._tables.get(0, []),
        }

        metadata = {
            "page_count": self.page_count,
            "has_tables": self.has_tables,
            "has_images": bool(doc.inline_shapes),
            "file_size": self.file_path.stat().st_size,
        }

        return content, metadata

    async def _load_xlsx(self) -> tuple[dict, dict]:
        """Load an Excel document."""
        wb = load_workbook(str(self.file_path), data_only=True)

        self.has_tables = True
        sheets_data = {}

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []
            for row in sheet.iter_rows(values_only=True):
                rows.append([str(cell) if cell else "" for cell in row])
            sheets_data[sheet_name] = rows

        self.page_count = len(wb.sheetnames)

        content = {
            "full_text": "",
            "sheets": sheets_data,
            "pages": [str(sheets_data)],
        }

        metadata = {
            "page_count": self.page_count,
            "has_tables": True,
            "has_images": False,
            "sheet_names": wb.sheetnames,
            "file_size": self.file_path.stat().st_size,
        }

        return content, metadata

    async def _load_csv(self) -> tuple[dict, dict]:
        """Load a CSV document."""
        import csv

        rows = []
        with open(self.file_path, newline="", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(row)

        self.has_tables = True
        self.page_count = 1
        self._tables[0] = [{"rows": rows}]

        content = {
            "full_text": "\n".join([",".join(row) for row in rows]),
            "tables": [{"headers": rows[0] if rows else [], "rows": rows[1:]}],
            "pages": [str(rows)],
        }

        metadata = {
            "page_count": 1,
            "has_tables": True,
            "has_images": False,
            "row_count": len(rows),
            "file_size": self.file_path.stat().st_size,
        }

        return content, metadata

    async def read_pages(self, page_numbers: list[int]) -> str:
        """Read text from specific pages (1-indexed)."""
        result = []
        for page_num in page_numbers:
            if 1 <= page_num <= len(self._pages):
                result.append(f"--- Page {page_num} ---\n{self._pages[page_num - 1]}")
        return "\n\n".join(result)

    async def find_section(
        self, section_hint: str, page_numbers: Optional[list[int]] = None
    ) -> str:
        """Find and return content from a section matching the hint."""
        hint_lower = section_hint.lower()

        for section in self._sections:
            if hint_lower in section.title.lower():
                if page_numbers is None or section.page_number in page_numbers:
                    return f"Section: {section.title}\n\n{section.content}"

        # Fallback: search in page text
        if page_numbers:
            return await self.read_pages(page_numbers)

        return f"Section '{section_hint}' not found"

    async def get_sections(self) -> list[dict]:
        """Get list of document sections."""
        return [
            {"title": s.title, "page": s.page_number, "level": s.level}
            for s in self._sections
        ]

    async def search(self, query: str) -> list[SearchMatch]:
        """Search document for text matching query."""
        matches = []
        pattern = re.compile(re.escape(query), re.IGNORECASE)

        for page_num, page_text in enumerate(self._pages):
            for match in pattern.finditer(page_text):
                matches.append(
                    SearchMatch(
                        text=match.group(),
                        page=page_num + 1,
                        position=match.start(),
                    )
                )

        return matches

    async def get_context(
        self, page: int, position: int, context_chars: int = 200
    ) -> str:
        """Get text context around a position."""
        if not (1 <= page <= len(self._pages)):
            return ""

        text = self._pages[page - 1]
        start = max(0, position - context_chars)
        end = min(len(text), position + context_chars)

        return text[start:end]

    async def get_page_image(self, page_number: int) -> Image.Image:
        """Get page as image for OCR."""
        if self.document_type == "image":
            return Image.open(self.file_path)

        if self.document_type == "pdf":
            doc = fitz.open(str(self.file_path))
            page = doc[page_number - 1]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for quality
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            doc.close()
            return img

        raise ValueError(f"Cannot get page image for {self.document_type}")

    async def get_images(self, page_number: int) -> list[bytes]:
        """Get images from a specific page."""
        return self._images.get(page_number - 1, [])


async def load_document(
    file_path: str, document_type: Literal["pdf", "image", "docx", "xlsx", "csv"]
) -> tuple[dict, dict]:
    """Load a document and return content and metadata.

    Also sets the parser as current context for tools to access.
    """
    parser = DocumentParser(
        file_path=Path(file_path),
        document_type=document_type,
    )

    content, metadata = await parser.load()
    parser.set_current()

    return content, metadata
